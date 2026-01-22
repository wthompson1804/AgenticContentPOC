"""
Export module for generating professional documentation.

This module provides export capabilities for:
- Markdown reports
- PDF documents
- DOCX documents
- HTML visualizations
- Complete assessment packages
"""

import io
import re
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_markdown_report(
    form_data: Dict[str, Any],
    research_results: Dict[str, Any],
    requirements_output: Dict[str, Any],
    agent_design_output: Dict[str, Any],
    capability_mapping: Dict[str, Any]
) -> str:
    """Generate a complete markdown assessment report.

    Args:
        form_data: User input form data
        research_results: Research findings
        requirements_output: Generated requirements
        agent_design_output: Agent design assessment
        capability_mapping: Capability mappings

    Returns:
        Complete markdown report string
    """
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

    agent_type = "N/A"
    if agent_design_output:
        agent_type = agent_design_output.get(
            'confirmed_type',
            agent_design_output.get('recommended_type', 'N/A')
        )

    report = f"""# DTC AI Agent Capability Assessment Report

**Generated:** {timestamp}
**Methodology:** Digital Twin Consortium AI Agent Capabilities Periodic Table (CPT)

---

## Executive Summary

| Parameter | Value |
|-----------|-------|
| Industry | {form_data.get('industry', 'N/A')} |
| Jurisdiction | {form_data.get('jurisdiction', 'N/A')} |
| Agent Type | {agent_type} |
| Capabilities Mapped | {capability_mapping.get('total_mapped', 0) if capability_mapping else 0} |
| Essential Capabilities | {capability_mapping.get('essential_count', 0) if capability_mapping else 0} |

---

## 1. Use Case Definition

### Industry Context
**Industry:** {form_data.get('industry', 'N/A')}
**Jurisdiction:** {form_data.get('jurisdiction', 'N/A')}
**Organization Size:** {form_data.get('organization_size', 'N/A')}
**Timeline:** {form_data.get('timeline', 'N/A')}

### Use Case Description
{form_data.get('use_case', 'No use case provided')}

### Existing Systems
{form_data.get('existing_systems', 'None specified')}

### Safety Requirements
{form_data.get('safety_requirements', 'None specified')}

---

## 2. Research Findings

"""

    if research_results:
        preliminary = research_results.get('preliminary_assessment', {})
        report += f"""### Preliminary Assessment
- **Go/No-Go Recommendation:** {preliminary.get('go_no_go', 'N/A').upper()}
- **Recommended Agent Type:** {preliminary.get('recommended_type', 'N/A')}
- **Confidence Level:** {preliminary.get('confidence_level', 'N/A').upper()}

### Key Risk Factors
"""
        for risk in preliminary.get('key_risks', ['None identified']):
            report += f"- {risk}\n"

        report += "\n### Research Areas\n\n"

        areas = research_results.get('research_areas', {})
        for area_key, area_data in areas.items():
            area_name = area_data.get('name', area_key)
            findings = area_data.get('findings', 'No findings')
            confidence = area_data.get('confidence', 'medium')
            report += f"#### {area_name}\n**Confidence:** {confidence.upper()}\n\n{findings}\n\n"
    else:
        report += "*Research not conducted*\n\n"

    report += """---

## 3. Business Requirements

"""

    if requirements_output:
        report += requirements_output.get('full_text', '*Requirements not generated*')
    else:
        report += "*Requirements not generated*"

    report += f"""

---

## 4. Agent Design

### Recommended Agent Type: {agent_type}

"""

    if agent_design_output:
        type_info = agent_design_output.get('type_info', {})
        report += f"""**Type Name:** {type_info.get('name', 'N/A')}
**Description:** {type_info.get('description', 'N/A')}

### Justification
{agent_design_output.get('justification', 'No justification provided')}

### Architecture Summary
{agent_design_output.get('architecture_summary', 'No architecture summary')}

### Full Design Document
{agent_design_output.get('full_document', '*Design document not available*')}
"""
    else:
        report += "*Agent design not generated*\n"

    report += """
---

## 5. Capability Mapping

"""

    if capability_mapping:
        report += f"""### Summary
- **Total Capabilities Mapped:** {capability_mapping.get('total_mapped', 0)}
- **Essential:** {capability_mapping.get('essential_count', 0)}
- **Advanced:** {capability_mapping.get('advanced_count', 0)}
- **Optional:** {capability_mapping.get('optional_count', 0)}

### Essential Capabilities
"""
        for cap_id in capability_mapping.get('essential_capabilities', []):
            report += f"- {cap_id}\n"

        report += "\n### Advanced Capabilities\n"
        for cap_id in capability_mapping.get('advanced_capabilities', []):
            report += f"- {cap_id}\n"

        report += "\n### Optional Capabilities\n"
        for cap_id in capability_mapping.get('optional_capabilities', []):
            report += f"- {cap_id}\n"

        report += "\n### Detailed Mapping\n"
        report += capability_mapping.get('full_document', '*Mapping document not available*')
    else:
        report += "*Capability mapping not generated*\n"

    report += f"""

---

## Appendix

### Methodology Reference
This assessment follows the Digital Twin Consortium's AI Agent Capabilities Periodic Table (CPT) framework, which organizes 45 capabilities across 6 categories:

1. **PK - Perception & Knowledge:** Environmental awareness and knowledge access
2. **CG - Cognition & Reasoning:** Planning, reasoning, and decision-making
3. **LA - Learning & Adaptation:** Memory, learning, and self-optimization
4. **AE - Action & Execution:** Task execution and tool integration
5. **IC - Interaction & Collaboration:** Communication and coordination
6. **GS - Governance & Safety:** Deployment, monitoring, and compliance

### Agent Types (T0-T4)
- **T0:** Static Automation - Rule-based, no learning
- **T1:** Conversational Agents - NLP interaction, basic context
- **T2:** Procedural Workflow Agents - Multi-step execution, tool integration
- **T3:** Cognitive Autonomous Agents - Self-directed planning, learning
- **T4:** Multi-Agent Generative Systems (MAGS) - Collaborative intelligence

---

*Report generated by DTC AI Agent Capability Assessment Tool*
*Powered by Digital Twin Consortium CPT Framework and Anthropic Claude*
"""

    return report


def generate_executive_summary(
    form_data: Dict[str, Any],
    research_results: Dict[str, Any],
    agent_design_output: Dict[str, Any],
    capability_mapping: Dict[str, Any]
) -> str:
    """Generate a brief executive summary.

    Args:
        form_data: User input form data
        research_results: Research findings
        agent_design_output: Agent design assessment
        capability_mapping: Capability mappings

    Returns:
        Executive summary string
    """
    agent_type = "N/A"
    if agent_design_output:
        agent_type = agent_design_output.get(
            'confirmed_type',
            agent_design_output.get('recommended_type', 'N/A')
        )

    go_no_go = "N/A"
    if research_results:
        go_no_go = research_results.get('preliminary_assessment', {}).get('go_no_go', 'N/A')

    summary = f"""# Executive Summary

## AI Agent Capability Assessment

**Industry:** {form_data.get('industry', 'N/A')}
**Use Case:** {form_data.get('use_case', 'N/A')[:200]}...

### Key Findings

| Metric | Value |
|--------|-------|
| Go/No-Go | {go_no_go.upper()} |
| Recommended Agent Type | {agent_type} |
| Capabilities Required | {capability_mapping.get('total_mapped', 0) if capability_mapping else 0} |
| Essential Capabilities | {capability_mapping.get('essential_count', 0) if capability_mapping else 0} |

### Recommendation

Based on the assessment, we recommend proceeding with a **{agent_type}** agent architecture.

"""

    if agent_design_output and agent_design_output.get('justification'):
        summary += f"**Justification:** {agent_design_output['justification'][:500]}"

    return summary


def export_to_html_package(
    form_data: Dict[str, Any],
    research_results: Dict[str, Any],
    requirements_output: Dict[str, Any],
    agent_design_output: Dict[str, Any],
    capability_mapping: Dict[str, Any]
) -> str:
    """Generate a self-contained HTML package with all assessment data.

    Args:
        form_data: User input form data
        research_results: Research findings
        requirements_output: Generated requirements
        agent_design_output: Agent design assessment
        capability_mapping: Capability mappings

    Returns:
        Self-contained HTML string
    """
    markdown_report = generate_markdown_report(
        form_data,
        research_results,
        requirements_output,
        agent_design_output,
        capability_mapping
    )

    # Convert markdown to basic HTML
    html_content = markdown_report.replace('\n', '<br>\n')
    html_content = html_content.replace('# ', '<h1>').replace('\n<br>', '</h1>\n')
    html_content = html_content.replace('## ', '<h2>').replace('\n<br>', '</h2>\n')
    html_content = html_content.replace('### ', '<h3>').replace('\n<br>', '</h3>\n')

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>DTC AI Agent Capability Assessment</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.6;
            color: #1F2937;
        }}
        h1 {{ color: #1F2937; border-bottom: 3px solid #3B82F6; padding-bottom: 10px; }}
        h2 {{ color: #374151; margin-top: 30px; }}
        h3 {{ color: #4B5563; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #E5E7EB; padding: 12px; text-align: left; }}
        th {{ background: #F3F4F6; }}
        hr {{ border: none; border-top: 1px solid #E5E7EB; margin: 30px 0; }}
        .header {{ text-align: center; margin-bottom: 40px; }}
        .footer {{ text-align: center; margin-top: 40px; color: #6B7280; font-size: 0.9rem; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>DTC AI Agent Capability Assessment</h1>
        <p>Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
    </div>

    <div class="content">
        {html_content}
    </div>

    <div class="footer">
        <p>Powered by Digital Twin Consortium CPT Framework and Anthropic Claude</p>
    </div>
</body>
</html>"""

    return html


def generate_docx_report(
    form_data: Dict[str, Any],
    research_results: Dict[str, Any],
    requirements_output: Dict[str, Any],
    agent_design_output: Dict[str, Any],
    capability_mapping: Dict[str, Any]
) -> Optional[bytes]:
    """Generate a DOCX report from the assessment data.

    Args:
        form_data: User input form data
        research_results: Research findings
        requirements_output: Generated requirements
        agent_design_output: Agent design assessment
        capability_mapping: Capability mappings

    Returns:
        DOCX file as bytes, or None if docx is not available
    """
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Title
    title = doc.add_heading('DTC AI Agent Capability Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {timestamp}\n").italic = True
    meta.add_run("Methodology: Digital Twin Consortium AI Agent CPT Framework").italic = True

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    agent_type = "N/A"
    if agent_design_output:
        agent_type = agent_design_output.get(
            'confirmed_type',
            agent_design_output.get('recommended_type', 'N/A')
        )

    # Summary table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('Industry', form_data.get('industry', 'N/A')),
        ('Jurisdiction', form_data.get('jurisdiction', 'N/A')),
        ('Agent Type', agent_type),
        ('Capabilities Mapped', str(capability_mapping.get('total_mapped', 0) if capability_mapping else 0)),
        ('Essential Capabilities', str(capability_mapping.get('essential_count', 0) if capability_mapping else 0)),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Use Case Definition
    doc.add_heading('1. Use Case Definition', level=1)

    doc.add_heading('Industry Context', level=2)
    p = doc.add_paragraph()
    p.add_run('Industry: ').bold = True
    p.add_run(f"{form_data.get('industry', 'N/A')}\n")
    p.add_run('Jurisdiction: ').bold = True
    p.add_run(f"{form_data.get('jurisdiction', 'N/A')}\n")
    p.add_run('Organization Size: ').bold = True
    p.add_run(f"{form_data.get('organization_size', 'N/A')}\n")
    p.add_run('Timeline: ').bold = True
    p.add_run(f"{form_data.get('timeline', 'N/A')}")

    doc.add_heading('Use Case Description', level=2)
    doc.add_paragraph(form_data.get('use_case', 'No use case provided'))

    if form_data.get('existing_systems'):
        doc.add_heading('Existing Systems', level=2)
        doc.add_paragraph(form_data.get('existing_systems', 'None specified'))

    if form_data.get('safety_requirements'):
        doc.add_heading('Safety Requirements', level=2)
        doc.add_paragraph(form_data.get('safety_requirements', 'None specified'))

    # Research Findings
    doc.add_heading('2. Research Findings', level=1)

    if research_results:
        preliminary = research_results.get('preliminary_assessment', {})

        doc.add_heading('Preliminary Assessment', level=2)
        p = doc.add_paragraph()
        p.add_run('Go/No-Go Recommendation: ').bold = True
        go_no_go = preliminary.get('go_no_go', 'N/A').upper()
        run = p.add_run(f"{go_no_go}\n")
        if go_no_go == 'GO':
            run.font.color.rgb = RGBColor(16, 185, 129)
        elif go_no_go == 'NO-GO':
            run.font.color.rgb = RGBColor(239, 68, 68)
        else:
            run.font.color.rgb = RGBColor(245, 158, 11)

        p.add_run('Recommended Agent Type: ').bold = True
        p.add_run(f"{preliminary.get('recommended_type', 'N/A')}\n")
        p.add_run('Confidence Level: ').bold = True
        p.add_run(f"{preliminary.get('confidence_level', 'N/A').upper()}")

        # Key Risks
        key_risks = preliminary.get('key_risks', [])
        if key_risks:
            doc.add_heading('Key Risk Factors', level=2)
            for risk in key_risks:
                doc.add_paragraph(risk, style='List Bullet')

        # Critical Success Factors
        success_factors = preliminary.get('critical_success_factors', [])
        if success_factors:
            doc.add_heading('Critical Success Factors', level=2)
            for factor in success_factors:
                doc.add_paragraph(factor, style='List Bullet')

        # Research Areas
        doc.add_heading('Research Areas', level=2)
        areas = research_results.get('research_areas', {})
        for area_key, area_data in areas.items():
            area_name = area_data.get('name', area_key)
            doc.add_heading(area_name, level=3)
            confidence = area_data.get('confidence', 'medium')
            p = doc.add_paragraph()
            p.add_run(f'Confidence: {confidence.upper()}').italic = True

            findings = area_data.get('findings', 'No findings')
            # Add findings, splitting into paragraphs
            for para in findings.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph('Research not conducted', style='Intense Quote')

    # Business Requirements
    doc.add_heading('3. Business Requirements', level=1)

    if requirements_output and requirements_output.get('full_text'):
        # Parse markdown to basic paragraphs
        req_text = requirements_output.get('full_text', '')
        for para in req_text.split('\n\n'):
            if para.strip():
                if para.startswith('#'):
                    # Convert markdown headers
                    header_level = len(para.split()[0])
                    header_text = para.lstrip('#').strip()
                    doc.add_heading(header_text, level=min(header_level + 1, 3))
                elif para.startswith('-') or para.startswith('*'):
                    # Convert bullet lists
                    for line in para.split('\n'):
                        if line.strip().startswith('-') or line.strip().startswith('*'):
                            doc.add_paragraph(line.strip().lstrip('-*').strip(), style='List Bullet')
                else:
                    doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph('Requirements not generated', style='Intense Quote')

    # Agent Design
    doc.add_heading('4. Agent Design', level=1)

    doc.add_heading(f'Recommended Agent Type: {agent_type}', level=2)

    if agent_design_output:
        type_info = agent_design_output.get('type_info', {})
        if type_info:
            p = doc.add_paragraph()
            p.add_run('Type Name: ').bold = True
            p.add_run(f"{type_info.get('name', 'N/A')}\n")
            p.add_run('Description: ').bold = True
            p.add_run(f"{type_info.get('description', 'N/A')}")

        if agent_design_output.get('justification'):
            doc.add_heading('Justification', level=3)
            doc.add_paragraph(agent_design_output['justification'])

        if agent_design_output.get('architecture_summary'):
            doc.add_heading('Architecture Summary', level=3)
            doc.add_paragraph(agent_design_output['architecture_summary'])
    else:
        doc.add_paragraph('Agent design not generated', style='Intense Quote')

    # Capability Mapping
    doc.add_heading('5. Capability Mapping', level=1)

    if capability_mapping:
        doc.add_heading('Summary', level=2)

        cap_table = doc.add_table(rows=4, cols=2)
        cap_table.style = 'Table Grid'
        cap_rows = [
            ('Total Capabilities Mapped', str(capability_mapping.get('total_mapped', 0))),
            ('Essential', str(capability_mapping.get('essential_count', 0))),
            ('Advanced', str(capability_mapping.get('advanced_count', 0))),
            ('Optional', str(capability_mapping.get('optional_count', 0))),
        ]
        for i, (label, value) in enumerate(cap_rows):
            cap_table.rows[i].cells[0].text = label
            cap_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Essential Capabilities
        essential = capability_mapping.get('essential_capabilities', [])
        if essential:
            doc.add_heading('Essential Capabilities', level=2)
            for cap_id in essential:
                doc.add_paragraph(cap_id, style='List Bullet')

        # Advanced Capabilities
        advanced = capability_mapping.get('advanced_capabilities', [])
        if advanced:
            doc.add_heading('Advanced Capabilities', level=2)
            for cap_id in advanced:
                doc.add_paragraph(cap_id, style='List Bullet')

        # Optional Capabilities
        optional = capability_mapping.get('optional_capabilities', [])
        if optional:
            doc.add_heading('Optional Capabilities', level=2)
            for cap_id in optional:
                doc.add_paragraph(cap_id, style='List Bullet')
    else:
        doc.add_paragraph('Capability mapping not generated', style='Intense Quote')

    # Appendix
    doc.add_heading('Appendix', level=1)

    doc.add_heading('Methodology Reference', level=2)
    doc.add_paragraph(
        'This assessment follows the Digital Twin Consortium\'s AI Agent Capabilities '
        'Periodic Table (CPT) framework, which organizes 45 capabilities across 6 categories:'
    )

    categories = [
        'PK - Perception & Knowledge: Environmental awareness and knowledge access',
        'CG - Cognition & Reasoning: Planning, reasoning, and decision-making',
        'LA - Learning & Adaptation: Memory, learning, and self-optimization',
        'AE - Action & Execution: Task execution and tool integration',
        'IC - Interaction & Collaboration: Communication and coordination',
        'GS - Governance & Safety: Deployment, monitoring, and compliance',
    ]
    for cat in categories:
        doc.add_paragraph(cat, style='List Number')

    doc.add_heading('Agent Types (T0-T4)', level=2)
    agent_types = [
        'T0: Static Automation - Rule-based, no learning',
        'T1: Conversational Agents - NLP interaction, basic context',
        'T2: Procedural Workflow Agents - Multi-step execution, tool integration',
        'T3: Cognitive Autonomous Agents - Self-directed planning, learning',
        'T4: Multi-Agent Generative Systems (MAGS) - Collaborative intelligence',
    ]
    for at in agent_types:
        doc.add_paragraph(at, style='List Bullet')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Report generated by DTC AI Agent Capability Assessment Tool\n').italic = True
    footer.add_run('Powered by Digital Twin Consortium CPT Framework and Anthropic Claude').italic = True

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream.getvalue()


# =============================================================================
# Phase 5: Additional Export Formats (PRD Part 17)
# =============================================================================

def render_internal_brief_md(
    artifact_doc: Dict[str, Any],
    form_data: Dict[str, Any],
    assumptions: Optional[list] = None
) -> str:
    """Generate internal team brief in markdown format.

    This format is designed for internal stakeholders and includes
    detailed assumptions and technical context.

    Args:
        artifact_doc: The 2-pager artifact document
        form_data: User input form data
        assumptions: Optional list of assumptions from chat intake

    Returns:
        Markdown string for internal brief
    """
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Extract sections from artifact
    sections = artifact_doc.get('sections', {})

    brief = f"""# Internal Team Brief: AI Agent Assessment

**Generated:** {timestamp}
**Status:** Internal Use Only

---

## Overview

**Industry:** {form_data.get('industry', 'N/A')}
**Jurisdiction:** {form_data.get('jurisdiction', 'N/A')}

### Problem Statement
{sections.get('problem_statement', {}).get('content', 'Not yet defined')}

### Proposed Solution
{sections.get('proposed_solution', {}).get('content', 'Not yet defined')}

---

## Technical Details

### Agent Type Recommendation
{sections.get('agent_type', {}).get('content', 'Not yet determined')}

### Key Capabilities Required
{sections.get('key_capabilities', {}).get('content', 'Not yet mapped')}

---

## Implementation Context

### Boundaries & Constraints
{sections.get('boundaries', {}).get('content', 'No explicit boundaries defined')}

### Success Metrics
{sections.get('success_metrics', {}).get('content', 'Not yet defined')}

---

## Risk Assessment

### Identified Risks
{sections.get('risks', {}).get('content', 'Risk assessment pending')}

### Mitigation Strategies
{sections.get('mitigations', {}).get('content', 'Mitigation strategies pending')}

"""

    # Add assumptions section if available
    if assumptions:
        brief += """---

## Working Assumptions

The following assumptions were identified during the intake process:

"""
        for i, assumption in enumerate(assumptions, 1):
            status = assumption.get('status', 'assumed')
            impact = assumption.get('impact', 'medium')
            confidence = assumption.get('confidence', 'med')
            statement = assumption.get('statement', '')

            status_emoji = "✓" if status == "confirmed" else "?" if status == "assumed" else "✗"
            brief += f"{i}. {status_emoji} **{statement}**\n"
            brief += f"   - Impact: {impact.upper()} | Confidence: {confidence.upper()} | Status: {status}\n\n"

    brief += f"""---

## Next Steps

{sections.get('next_steps', {}).get('content', 'Next steps to be determined')}

---

*This internal brief was generated from the AI Agent Assessment Tool.*
*For external distribution, use the Executive Brief format.*
"""

    return brief


def render_exec_brief_md(
    artifact_doc: Dict[str, Any],
    form_data: Dict[str, Any]
) -> str:
    """Generate executive brief in markdown format.

    This format is designed for executive stakeholders and focuses
    on business value, high-level recommendations, and key decisions.

    Args:
        artifact_doc: The 2-pager artifact document
        form_data: User input form data

    Returns:
        Markdown string for executive brief
    """
    timestamp = datetime.now().strftime("%Y-%m-%d")

    sections = artifact_doc.get('sections', {})

    # Extract agent type for summary
    agent_type_content = sections.get('agent_type', {}).get('content', '')
    agent_type = "T2"  # Default
    if 'T0' in agent_type_content:
        agent_type = "T0"
    elif 'T1' in agent_type_content:
        agent_type = "T1"
    elif 'T3' in agent_type_content:
        agent_type = "T3"
    elif 'T4' in agent_type_content:
        agent_type = "T4"

    brief = f"""# Executive Brief: AI Agent Initiative

**Date:** {timestamp}
**Industry:** {form_data.get('industry', 'N/A')}

---

## Executive Summary

**Recommendation:** Proceed with **{agent_type}** agent implementation

### The Opportunity
{sections.get('problem_statement', {}).get('content', 'Opportunity assessment in progress')[:500]}

### Proposed Approach
{sections.get('proposed_solution', {}).get('content', 'Solution design in progress')[:500]}

---

## Key Findings

### Agent Architecture
{sections.get('agent_type', {}).get('content', 'Architecture recommendation pending')[:300]}

### Critical Success Factors
{sections.get('success_metrics', {}).get('content', 'Success metrics to be defined')[:300]}

---

## Risk Summary

{sections.get('risks', {}).get('content', 'Risk assessment in progress')[:400]}

---

## Investment Considerations

| Aspect | Assessment |
|--------|------------|
| Complexity | {'High' if agent_type in ['T3', 'T4'] else 'Moderate' if agent_type == 'T2' else 'Low'} |
| Timeline | {form_data.get('timeline', 'To be determined')} |
| Agent Type | {agent_type} |

---

## Recommended Next Steps

{sections.get('next_steps', {}).get('content', 'Next steps pending assessment completion')[:400]}

---

*Executive Brief generated {timestamp}*
"""

    return brief


def render_email_md(
    artifact_doc: Dict[str, Any],
    form_data: Dict[str, Any],
    recipient_type: str = "stakeholder"
) -> str:
    """Generate email-ready summary in markdown format.

    Args:
        artifact_doc: The 2-pager artifact document
        form_data: User input form data
        recipient_type: Type of recipient (stakeholder, technical, executive)

    Returns:
        Markdown string formatted as email body
    """
    sections = artifact_doc.get('sections', {})

    # Determine appropriate tone based on recipient
    if recipient_type == "executive":
        greeting = "Dear Leadership Team,"
        closing = "Best regards,"
    elif recipient_type == "technical":
        greeting = "Hi Team,"
        closing = "Thanks,"
    else:
        greeting = "Hello,"
        closing = "Best,"

    email = f"""{greeting}

I wanted to share a summary of the AI agent capability assessment we recently completed for our {form_data.get('industry', '')} initiative.

**Overview**

{sections.get('problem_statement', {}).get('content', 'Assessment overview pending')[:300]}

**Recommendation**

{sections.get('proposed_solution', {}).get('content', 'Recommendations being finalized')[:300]}

**Key Points**

- Agent Type: {sections.get('agent_type', {}).get('content', 'TBD')[:100]}
- Key Capabilities: {sections.get('key_capabilities', {}).get('content', 'Being mapped')[:100]}

**Next Steps**

{sections.get('next_steps', {}).get('content', 'To be discussed')[:200]}

I'd be happy to schedule time to discuss this in more detail.

{closing}


---
*This summary was generated from the DTC AI Agent Assessment Tool*
"""

    return email


def render_slide_outline_md(
    artifact_doc: Dict[str, Any],
    form_data: Dict[str, Any]
) -> str:
    """Generate presentation slide outline in markdown format.

    Args:
        artifact_doc: The 2-pager artifact document
        form_data: User input form data

    Returns:
        Markdown string with slide-by-slide outline
    """
    sections = artifact_doc.get('sections', {})

    outline = f"""# Presentation Outline: AI Agent Capability Assessment

**Industry:** {form_data.get('industry', 'N/A')}

---

## Slide 1: Title

**AI Agent Capability Assessment**
- Industry: {form_data.get('industry', 'N/A')}
- Date: {datetime.now().strftime("%B %Y")}
- Methodology: DTC Capabilities Periodic Table (CPT)

---

## Slide 2: The Challenge

**Problem Statement**

{sections.get('problem_statement', {}).get('content', 'Challenge description')[:400]}

*Speaker Notes: Emphasize business impact and urgency*

---

## Slide 3: Proposed Solution

**Our Approach**

{sections.get('proposed_solution', {}).get('content', 'Solution overview')[:400]}

*Speaker Notes: Focus on differentiation and value proposition*

---

## Slide 4: Agent Architecture

**Recommended: Agent Type**

{sections.get('agent_type', {}).get('content', 'Architecture details')[:300]}

**Visual:** Agent type spectrum diagram (T0-T4)

*Speaker Notes: Explain why this type was selected*

---

## Slide 5: Key Capabilities

**Required Capabilities**

{sections.get('key_capabilities', {}).get('content', 'Capability mapping')[:400]}

**Visual:** CPT periodic table with highlighted capabilities

---

## Slide 6: Success Metrics

**How We'll Measure Success**

{sections.get('success_metrics', {}).get('content', 'Metrics overview')[:300]}

*Speaker Notes: Tie metrics to business outcomes*

---

## Slide 7: Risk Assessment

**Key Risks & Mitigations**

Risks:
{sections.get('risks', {}).get('content', 'Risk list')[:200]}

Mitigations:
{sections.get('mitigations', {}).get('content', 'Mitigation strategies')[:200]}

---

## Slide 8: Implementation Roadmap

**Next Steps**

{sections.get('next_steps', {}).get('content', 'Implementation plan')[:300]}

**Visual:** Timeline or Gantt chart

---

## Slide 9: Boundaries & Scope

**What This Initiative Will NOT Do**

{sections.get('boundaries', {}).get('content', 'Out of scope items')[:300]}

*Speaker Notes: Setting clear expectations*

---

## Slide 10: Q&A / Discussion

**Questions?**

- Contact: [Your Name]
- Assessment Tool: DTC AI Agent CPT Framework

---

*Slide outline generated from AI Agent Assessment Tool*
"""

    return outline


def get_available_export_formats() -> list:
    """Return list of available export formats for UI display.

    Returns:
        List of format dictionaries with id, name, description
    """
    formats = [
        {
            "id": "markdown",
            "name": "Full Markdown Report",
            "description": "Complete assessment report in Markdown format",
            "function": "generate_markdown_report"
        },
        {
            "id": "html",
            "name": "HTML Package",
            "description": "Self-contained HTML with styling",
            "function": "export_to_html_package"
        },
        {
            "id": "exec_summary",
            "name": "Executive Summary",
            "description": "Brief summary for leadership",
            "function": "generate_executive_summary"
        },
        {
            "id": "internal_brief",
            "name": "Internal Team Brief",
            "description": "Detailed brief for internal stakeholders",
            "function": "render_internal_brief_md"
        },
        {
            "id": "exec_brief",
            "name": "Executive Brief",
            "description": "High-level brief for executives",
            "function": "render_exec_brief_md"
        },
        {
            "id": "email",
            "name": "Email Summary",
            "description": "Email-ready summary",
            "function": "render_email_md"
        },
        {
            "id": "slides",
            "name": "Slide Outline",
            "description": "Presentation slide outline",
            "function": "render_slide_outline_md"
        },
    ]

    if DOCX_AVAILABLE:
        formats.insert(2, {
            "id": "docx",
            "name": "Word Document",
            "description": "Professional DOCX report",
            "function": "generate_docx_report"
        })

    return formats
