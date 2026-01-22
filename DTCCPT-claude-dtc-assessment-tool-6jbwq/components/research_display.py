"""
Research display component for showing research results with citations.
"""

import io
import streamlit as st
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_research_docx(results: Dict[str, Any]) -> Optional[bytes]:
    """Generate a DOCX file from research results.

    Args:
        results: Research results dictionary

    Returns:
        DOCX file as bytes, or None if docx not available
    """
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Title
    title = doc.add_heading('DTC AI Agent Research Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {timestamp}").italic = True

    doc.add_paragraph()

    # Summary section
    summary = results.get('summary', {})
    doc.add_heading('Assessment Context', level=1)
    p = doc.add_paragraph()
    p.add_run('Industry: ').bold = True
    p.add_run(f"{summary.get('industry', 'N/A')}\n")
    p.add_run('Use Case: ').bold = True
    p.add_run(f"{summary.get('use_case', 'N/A')}\n")
    p.add_run('Jurisdiction: ').bold = True
    p.add_run(f"{summary.get('jurisdiction', 'N/A')}")

    # Preliminary Assessment
    assessment = results.get('preliminary_assessment', {})
    doc.add_heading('Preliminary Assessment', level=1)

    p = doc.add_paragraph()
    p.add_run('Go/No-Go Recommendation: ').bold = True
    p.add_run(f"{assessment.get('go_no_go', 'N/A').upper()}\n")
    p.add_run('Recommended Agent Type: ').bold = True
    p.add_run(f"{assessment.get('recommended_type', 'N/A')}\n")
    p.add_run('Confidence Level: ').bold = True
    p.add_run(f"{assessment.get('confidence_level', 'N/A').upper()}")

    # Key Risks
    key_risks = assessment.get('key_risks', [])
    if key_risks:
        doc.add_heading('Key Risk Factors', level=2)
        for risk in key_risks:
            doc.add_paragraph(risk, style='List Bullet')

    # Success Factors
    success_factors = assessment.get('critical_success_factors', [])
    if success_factors:
        doc.add_heading('Critical Success Factors', level=2)
        for factor in success_factors:
            doc.add_paragraph(factor, style='List Bullet')

    # Rationale
    rationale = assessment.get('recommendation_rationale', '')
    if rationale:
        doc.add_heading('Recommendation Rationale', level=2)
        doc.add_paragraph(rationale)

    # Research Areas
    doc.add_heading('Research Findings', level=1)
    areas = results.get('research_areas', {})
    for area_key, area_data in areas.items():
        area_name = area_data.get('name', area_key)
        doc.add_heading(area_name, level=2)

        confidence = area_data.get('confidence', 'medium')
        p = doc.add_paragraph()
        p.add_run(f'Confidence: {confidence.upper()}').italic = True

        # Add summary if available
        summary_text = area_data.get('summary', '')
        if summary_text:
            doc.add_paragraph(summary_text)

        # Add full findings
        findings = area_data.get('findings', '')
        if findings:
            for para in findings.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())

    # Full content at the end
    if results.get('full_content'):
        doc.add_heading('Complete Research Report', level=1)
        for para in results['full_content'].split('\n\n'):
            if para.strip():
                if para.startswith('#'):
                    header_text = para.lstrip('#').strip()
                    level = min(len(para.split()[0]) if para.split() else 1, 3)
                    doc.add_heading(header_text, level=level)
                else:
                    doc.add_paragraph(para.strip())

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Powered by Digital Twin Consortium CPT Framework').italic = True

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream.getvalue()


# Educational descriptions for agent types
AGENT_TYPE_INFO = {
    "T0": {
        "name": "Static Automation",
        "short": "Rule-based, deterministic systems",
        "description": """**What it is:** Pre-programmed automation that follows fixed rules without learning or adaptation.

**Best for:** Simple, repetitive tasks with predictable outcomes where the rules are well-defined.

**Examples:** Scheduled reports, threshold-based alerts, simple data transformations.

**Key characteristics:**
- No learning capability
- Fully deterministic behavior
- Requires manual updates for changes
- Lowest complexity and risk""",
    },
    "T1": {
        "name": "Conversational Agent",
        "short": "Natural language interaction with basic context",
        "description": """**What it is:** AI that can understand and respond to natural language, maintaining basic conversation context.

**Best for:** Customer service, information retrieval, simple Q&A interfaces.

**Examples:** Chatbots, voice assistants, FAQ systems.

**Key characteristics:**
- Natural language understanding
- Session-based context memory
- Limited reasoning capability
- Human-like interaction pattern""",
    },
    "T2": {
        "name": "Procedural Workflow Agent",
        "short": "Multi-step task execution with tool integration",
        "description": """**What it is:** AI that can execute multi-step workflows, use external tools, and coordinate between systems.

**Best for:** Process automation, data pipeline orchestration, system integration tasks.

**Examples:** Automated data processing, multi-system workflows, report generation.

**Key characteristics:**
- Multi-step task execution
- Tool and API integration
- Error handling and recovery
- Structured workflow management""",
    },
    "T3": {
        "name": "Cognitive Autonomous Agent",
        "short": "Self-directed planning with learning and adaptation",
        "description": """**What it is:** AI that can plan its own actions, learn from experience, and adapt to changing conditions autonomously.

**Best for:** Complex decision-making, dynamic environments, situations requiring judgment and learning.

**Examples:** Predictive maintenance, autonomous optimization, adaptive control systems.

**Key characteristics:**
- Self-directed goal planning
- Learning from outcomes
- Adaptive behavior
- Requires careful governance and oversight""",
    },
    "T4": {
        "name": "Multi-Agent System (MAGS)",
        "short": "Collaborative AI with distributed coordination",
        "description": """**What it is:** Multiple AI agents working together, each with specialized roles, coordinating to achieve complex objectives.

**Best for:** Large-scale complex systems, scenarios requiring multiple specialized capabilities working in concert.

**Examples:** Smart grid management, autonomous fleet coordination, complex simulation systems.

**Key characteristics:**
- Multiple specialized agents
- Inter-agent communication
- Emergent collective behavior
- Highest complexity and capability""",
    },
}


def render_research_results(results: Dict[str, Any]) -> None:
    """Render the complete research results.

    Args:
        results: Formatted research results from format_research_for_display()
    """
    if results.get("error"):
        st.error(f"Research Error: {results['error']}")
        return

    # Preliminary Assessment Card
    render_preliminary_assessment(results.get("preliminary_assessment", {}))

    st.divider()

    # Research Areas
    st.markdown("### Research Findings")
    st.caption("Click each tab to view detailed findings. Expand sections for more detail.")

    research_areas = results.get("research_areas", {})
    full_content = results.get("full_content", "")

    # Check if any sections have content
    has_any_content = any(
        area.get("findings") or area.get("summary")
        for area in research_areas.values()
    )

    # If no sections extracted but we have full_content, show a warning
    if not has_any_content and full_content:
        st.warning("Section extraction encountered issues. Content is available in the 'Full Report' section below and in each tab.")

    # Create tabs for each research area
    tabs = st.tabs([
        "Industry Adoption",
        "Regulatory",
        "Technical",
        "Risk & Failure",
        "Economic"
    ])

    with tabs[0]:
        render_research_area(
            research_areas.get("industry_adoption", {}),
            "Industry AI Adoption",
            "Current state of AI deployment in this industry, success rates, and adoption patterns.",
            full_content
        )

    with tabs[1]:
        render_research_area(
            research_areas.get("regulatory_environment", {}),
            "Regulatory Environment",
            "Relevant regulations, standards, and compliance requirements for AI systems.",
            full_content
        )

    with tabs[2]:
        render_research_area(
            research_areas.get("technical_integration", {}),
            "Technical Integration",
            "Common technology stacks, integration patterns, and technical challenges.",
            full_content
        )

    with tabs[3]:
        render_research_area(
            research_areas.get("risk_failure_modes", {}),
            "Risk & Failure Modes",
            "Documented failures, root causes, and risk factors to consider.",
            full_content
        )

    with tabs[4]:
        render_research_area(
            research_areas.get("economic_viability", {}),
            "Economic Viability",
            "ROI expectations, cost structures, and economic considerations.",
            full_content
        )

    # Full research report section
    if results.get("full_content"):
        st.divider()
        st.markdown("### üìÑ Full Research Report")

        st.caption("Download the complete research report for offline review or sharing.")

        col1, col2, col3 = st.columns([2, 1, 1])
        with col2:
            # Markdown download
            st.download_button(
                label="‚¨áÔ∏è Download .md",
                data=results["full_content"],
                file_name="dtc_research_report.md",
                mime="text/markdown",
                use_container_width=True
            )
        with col3:
            # DOCX download
            docx_bytes = generate_research_docx(results)
            if docx_bytes:
                st.download_button(
                    label="‚¨áÔ∏è Download .docx",
                    data=docx_bytes,
                    file_name="dtc_research_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        with st.expander("üìã Preview Full Report", expanded=False):
            st.markdown(results["full_content"])

    # Sources section
    if results.get("sources"):
        render_sources(results["sources"])


def render_preliminary_assessment(assessment: Dict[str, Any]) -> None:
    """Render the preliminary assessment card.

    Args:
        assessment: Preliminary assessment data
    """
    st.markdown("### Preliminary Assessment")

    # Traffic Light Indicator - Large and prominent
    go_no_go = assessment.get("go_no_go", "pending")
    go_config = {
        "go": {"color": "#10B981", "bg": "#10B98130", "icon": "‚óè", "label": "GO", "desc": "Research supports moving forward"},
        "caution": {"color": "#F59E0B", "bg": "#F59E0B30", "icon": "‚óè", "label": "CAUTION", "desc": "Proceed with careful planning"},
        "no-go": {"color": "#EF4444", "bg": "#EF444430", "icon": "‚óè", "label": "NO-GO", "desc": "Significant concerns identified"},
        "pending": {"color": "#6B7280", "bg": "#6B728030", "icon": "‚óã", "label": "PENDING", "desc": "Awaiting analysis"},
    }
    config = go_config.get(go_no_go, go_config["pending"])

    # Traffic light style indicator
    st.markdown(f"""
    <div style="
        display: flex;
        align-items: center;
        gap: 20px;
        background: linear-gradient(135deg, {config['bg']}, transparent);
        border-left: 6px solid {config['color']};
        border-radius: 8px;
        padding: 20px;
        margin-bottom: 20px;
    ">
        <div style="
            font-size: 4rem;
            color: {config['color']};
            line-height: 1;
            text-shadow: 0 0 20px {config['color']}40;
        ">{config['icon']}</div>
        <div>
            <div style="font-size: 1.5rem; font-weight: bold; color: {config['color']};">{config['label']}</div>
            <div style="color: #666;">{config['desc']}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Three columns for agent type and confidence
    col1, col2, col3 = st.columns([2, 2, 1])

    with col1:
        agent_type = assessment.get("recommended_type", "TBD")
        type_info = AGENT_TYPE_INFO.get(agent_type, {"name": "Unknown", "short": "Not determined"})

        st.markdown(f"""
        <div style="
            background-color: #3B82F615;
            border: 2px solid #3B82F6;
            border-radius: 8px;
            padding: 16px;
            text-align: center;
        ">
            <div style="font-size: 2rem; font-weight: bold; color: #3B82F6;">{agent_type}</div>
            <div style="font-weight: bold;">{type_info['name']}</div>
            <div style="font-size: 0.75rem; color: #666; margin-top: 4px;">{type_info['short']}</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        confidence = assessment.get("confidence_level", "medium")
        conf_config = {
            "high": {"color": "#10B981", "desc": "Strong evidence available"},
            "medium": {"color": "#F59E0B", "desc": "Moderate evidence"},
            "low": {"color": "#EF4444", "desc": "Limited information"},
        }
        conf = conf_config.get(confidence, conf_config["medium"])

        st.markdown(f"""
        <div style="
            background-color: {conf['color']}15;
            border: 2px solid {conf['color']};
            border-radius: 8px;
            padding: 16px;
            text-align: center;
        ">
            <div style="font-size: 1.5rem; font-weight: bold; color: {conf['color']}; text-transform: uppercase;">{confidence}</div>
            <div style="font-weight: bold;">Confidence</div>
            <div style="font-size: 0.75rem; color: #666; margin-top: 4px;">{conf['desc']}</div>
        </div>
        """, unsafe_allow_html=True)

    # Confidence improvement guidance (show when not high confidence)
    if confidence in ["low", "medium"]:
        with st.expander(f"How to improve confidence from {confidence.upper()} to HIGH", expanded=(confidence == "low")):
            st.markdown("""
            <div style="background-color: #FEF3C7; border-radius: 8px; padding: 16px; margin-bottom: 12px;">
                <strong style="color: #92400E;">To increase assessment confidence, provide:</strong>
            </div>
            """, unsafe_allow_html=True)

            improvement_items = [
                ("Specific use case details", "Describe the exact processes, decisions, or tasks the AI agent will handle"),
                ("Existing system inventory", "List current software, data sources, and integration points"),
                ("Quantified business metrics", "Include current performance baselines and target improvements (e.g., processing time, error rates)"),
                ("Regulatory requirements", "Specify any compliance frameworks, certifications, or audit requirements"),
                ("Risk tolerance statement", "Define acceptable levels of autonomy and failure modes"),
            ]

            for title, desc in improvement_items:
                st.markdown(f"""
                <div style="
                    display: flex;
                    align-items: flex-start;
                    gap: 10px;
                    margin-bottom: 10px;
                    padding: 10px;
                    background-color: #F9FAFB;
                    border-radius: 6px;
                    border-left: 3px solid #3B82F6;
                ">
                    <span style="color: #3B82F6; font-weight: bold;">*</span>
                    <div>
                        <div style="font-weight: 600; color: #1F2937;">{title}</div>
                        <div style="font-size: 0.85rem; color: #6B7280;">{desc}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            st.caption("Providing this information in the input form will enable more accurate research and higher confidence recommendations.")

    # Agent type explanation (expandable)
    if agent_type in AGENT_TYPE_INFO:
        with st.expander(f"‚ÑπÔ∏è What is a {agent_type} ({AGENT_TYPE_INFO[agent_type]['name']})?", expanded=False):
            st.markdown(AGENT_TYPE_INFO[agent_type]['description'])

    # Recommendation Rationale - THE WHY
    st.markdown("---")
    st.markdown("#### Why This Recommendation?")

    rationale = assessment.get("recommendation_rationale", "")
    if rationale:
        st.markdown(f"""
        <div style="
            background-color: #f8f9fa;
            border-radius: 8px;
            padding: 16px;
            margin-bottom: 16px;
            border-left: 4px solid {config['color']};
        ">
            {rationale}
        </div>
        """, unsafe_allow_html=True)
    else:
        st.info("Detailed rationale will be provided based on research findings.")

    # Key risks and success factors - with actual bullets
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("##### ‚ö†Ô∏è Key Risk Factors")
        key_risks = assessment.get("key_risks", [])
        if key_risks:
            for risk in key_risks:
                st.markdown(f"""
                <div style="
                    display: flex;
                    align-items: flex-start;
                    gap: 8px;
                    margin-bottom: 8px;
                    padding: 8px;
                    background-color: #FEF2F2;
                    border-radius: 4px;
                ">
                    <span style="color: #EF4444;">‚úó</span>
                    <span style="color: #374151;">{risk}</span>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.caption("_No specific risk factors extracted. Review detailed findings for risk information._")

    with col2:
        st.markdown("##### ‚úì Critical Success Factors")
        success_factors = assessment.get("critical_success_factors", [])
        if success_factors:
            for factor in success_factors:
                st.markdown(f"""
                <div style="
                    display: flex;
                    align-items: flex-start;
                    gap: 8px;
                    margin-bottom: 8px;
                    padding: 8px;
                    background-color: #F0FDF4;
                    border-radius: 4px;
                ">
                    <span style="color: #10B981;">‚úì</span>
                    <span style="color: #374151;">{factor}</span>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.caption("_No specific success factors extracted. Review detailed findings for guidance._")


def render_research_area(area: Dict[str, Any], title: str, description: str, full_content: str = "") -> None:
    """Render a single research area with summary and expandable full content.

    Args:
        area: Research area data
        title: Section title
        description: Section description
        full_content: Full research content to search as fallback
    """
    findings = area.get("findings", "")
    summary = area.get("summary", "")
    confidence = area.get("confidence", "medium")

    # Confidence indicator
    conf_colors = {"high": "#10B981", "medium": "#F59E0B", "low": "#EF4444"}
    conf_color = conf_colors.get(confidence, "#6B7280")

    st.markdown(f"""
    <div style="display: flex; align-items: center; margin-bottom: 12px; gap: 8px;">
        <span style="
            background-color: {conf_color};
            color: white;
            padding: 2px 8px;
            border-radius: 4px;
            font-size: 0.7rem;
            font-weight: bold;
        ">{confidence.upper()} CONFIDENCE</span>
    </div>
    """, unsafe_allow_html=True)

    st.caption(description)

    if summary or findings:
        # Display the summary prominently
        display_summary = summary if summary else (findings[:300] + "..." if len(findings) > 300 else findings)

        st.markdown(f"""
        <div style="
            background-color: #f8f9fa;
            border-radius: 8px;
            padding: 16px;
            margin: 12px 0;
            border-left: 3px solid {conf_color};
        ">
            <div style="font-weight: 500; margin-bottom: 8px; color: #374151;">Summary</div>
            <div style="color: #4B5563; line-height: 1.6;">{display_summary}</div>
        </div>
        """, unsafe_allow_html=True)

        # Show full content in expander if there's more to show
        if findings and len(findings) > len(display_summary) + 50:
            with st.expander("üìñ View Full Analysis", expanded=False):
                st.markdown(findings)
    else:
        # Try to find relevant content from full_content as fallback
        if full_content:
            fallback = _extract_section_fallback(full_content, title)
            if fallback:
                st.markdown(f"""
                <div style="
                    background-color: #FEF3C7;
                    border-radius: 8px;
                    padding: 16px;
                    margin: 12px 0;
                    border-left: 3px solid #F59E0B;
                ">
                    <div style="font-weight: 500; margin-bottom: 8px; color: #92400E;">Content (from full report)</div>
                    <div style="color: #4B5563; line-height: 1.6;">{fallback[:500]}{'...' if len(fallback) > 500 else ''}</div>
                </div>
                """, unsafe_allow_html=True)
                with st.expander("üìñ View Full Section", expanded=False):
                    st.markdown(fallback)
            else:
                st.info(f"No detailed findings available for {title}. Check the full report below.")
        else:
            st.info(f"No detailed findings available for {title}. This section will be populated with research results.")


def _extract_section_fallback(content: str, title: str) -> str:
    """Try to extract a section from content using simple patterns.

    Args:
        content: Full content to search
        title: Section title to find

    Returns:
        Extracted section content or empty string
    """
    import re

    # Simplify title for matching
    keywords = title.lower().split()

    # Try to find section by keywords
    lines = content.split('\n')
    in_section = False
    section_content = []

    for i, line in enumerate(lines):
        line_lower = line.lower()

        # Check if this line is a header containing our keywords
        is_header = line.strip().startswith('#') or (line.strip().startswith('**') and line.strip().endswith('**'))
        has_keywords = sum(1 for kw in keywords if kw in line_lower) >= min(2, len(keywords))

        if is_header and has_keywords:
            in_section = True
            continue
        elif in_section:
            # Check if we've hit the next section
            if line.strip().startswith('#') and len(line.strip().lstrip('#').strip()) > 0:
                break
            if line.strip().startswith('**') and line.strip().endswith('**') and len(line.strip()) > 4:
                # Could be a sub-header, check if it's a new section
                other_sections = ['industry', 'regulatory', 'technical', 'risk', 'economic', 'preliminary', 'executive', 'sources']
                if any(s in line_lower for s in other_sections if s not in title.lower()):
                    break
            section_content.append(line)

    return '\n'.join(section_content).strip()


def render_sources(sources: list) -> None:
    """Render the sources section.

    Args:
        sources: List of source dictionaries with 'title' and 'url' keys
    """
    with st.expander("üìö Sources & Citations", expanded=False):
        if sources:
            for source in sources:
                title = source.get("title", "Untitled")
                url = source.get("url", "")
                date = source.get("date", "")

                if url:
                    st.markdown(f"- [{title}]({url}) {f'({date})' if date else ''}")
                else:
                    st.markdown(f"- {title} {f'({date})' if date else ''}")
        else:
            st.caption("Sources will be listed here when available from research.")


def render_research_loading() -> None:
    """Render a loading state for research in progress."""
    st.markdown("### Conducting Deep Research")

    progress_text = st.empty()
    progress_bar = st.progress(0)

    research_steps = [
        "Analyzing industry AI adoption patterns...",
        "Reviewing regulatory environment...",
        "Evaluating technical integration requirements...",
        "Identifying risk and failure modes...",
        "Assessing economic viability...",
        "Synthesizing findings..."
    ]

    for i, step in enumerate(research_steps):
        progress_text.markdown(f"*{step}*")
        progress_bar.progress((i + 1) / len(research_steps))

    return progress_text, progress_bar


def render_research_error(error_message: str) -> None:
    """Render an error state for failed research.

    Args:
        error_message: Error message to display
    """
    st.error("Research Failed")

    st.markdown(f"""
    <div style="
        background-color: #FEE2E2;
        border: 1px solid #EF4444;
        border-radius: 8px;
        padding: 16px;
        margin: 16px 0;
    ">
        <strong>Error:</strong> {error_message}
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    **Possible causes:**
    - API key not configured or invalid
    - Network connectivity issues
    - Rate limiting

    **Suggested actions:**
    1. Check your Anthropic API key in the sidebar
    2. Verify network connectivity
    3. Try again in a few moments
    """)

    if st.button("Retry Research"):
        st.session_state.research_results = None
        st.rerun()
